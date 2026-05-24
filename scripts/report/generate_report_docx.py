"""
Generates Final_Project_Report.docx from the markdown source.
Run with: /mnt/d/EDITEASE/venv/bin/python generate_report_docx.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from PIL import Image

BASE_DIR = Path(__file__).resolve().parent
REPORT_PATH = BASE_DIR / "Final_Project_Report.md"
OUTPUT_PATH = BASE_DIR / "Final_Project_Report.docx"


def set_doc_margins(doc, top=2.54, bottom=2.54, left=3.0, right=2.54):
    """Set page margins in centimetres."""
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def configure_styles(doc):
    """Apply University-standard formatting to built-in styles."""
    styles = doc.styles

    # Normal (body text)
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = Pt(18)  # 1.5 lines

    # Heading 1 — Chapter title
    h1 = styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True

    # Heading 2 — Section
    h2 = styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.keep_with_next = True

    # Heading 3 — Subsection
    h3 = styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.italic = False
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.keep_with_next = True

    # List bullet
    try:
        lb = styles["List Bullet"]
        lb.font.name = "Times New Roman"
        lb.font.size = Pt(12)
        lb.paragraph_format.space_after = Pt(3)
    except Exception:
        pass

    # List number
    try:
        ln = styles["List Number"]
        ln.font.name = "Times New Roman"
        ln.font.size = Pt(12)
        ln.paragraph_format.space_after = Pt(3)
    except Exception:
        pass


def add_cover_page(doc):
    """Build a formatted cover page."""
    doc.add_paragraph()
    doc.add_paragraph()

    # University partner line
    uni_line = doc.add_paragraph()
    uni_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = uni_line.add_run("UNIVERSITY PARTNER")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.bold = True

    # University name
    uni_name = doc.add_paragraph()
    uni_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = uni_name.add_run("UNIVERSITY OF WOLVERHAMPTON")
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.font.bold = True

    college = doc.add_paragraph()
    college.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = college.add_run("Herald College Kathmandu")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)

    for _ in range(3):
        doc.add_paragraph()

    # Module title
    module = doc.add_paragraph()
    module.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = module.add_run("Project and Professionalism")
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(18)
    r3.font.bold = True

    for _ in range(2):
        doc.add_paragraph()

    # Project title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = title.add_run("EDITEASE: An AI-Assisted Video Organising Platform")
    r4.font.name = "Times New Roman"
    r4.font.size = Pt(20)
    r4.font.bold = True

    for _ in range(3):
        doc.add_paragraph()

    # Student details table
    details = [
        ("Name:", "Ashreen Dangol"),
        ("Student ID:", "2407774"),
        ("Supervisor:", "Kaushal Kishor Mishra"),
        ("University Partner:", "University of Wolverhampton"),
        ("Institution:", "Herald College Kathmandu"),
        ("Submission Date:", "[TO FILL: Final submission date]"),
    ]
    for label, value in details:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_label = p.add_run(label + "  ")
        run_label.font.name = "Times New Roman"
        run_label.font.size = Pt(12)
        run_label.font.bold = True
        run_value = p.add_run(value)
        run_value.font.name = "Times New Roman"
        run_value.font.size = Pt(12)

    doc.add_page_break()


def add_table_from_md(doc, lines, start_idx):
    """Parse a markdown table starting at start_idx and add a Word table. Returns last index used."""
    table_lines = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i].strip())
        i += 1

    if len(table_lines) < 2:
        return start_idx

    # Parse header
    headers = [c.strip() for c in table_lines[0].split("|") if c.strip()]
    # Skip separator line (contains ---)
    data_rows = []
    for row_line in table_lines[2:]:
        cells = [c.strip() for c in row_line.split("|") if c.strip() != ""]
        if cells:
            data_rows.append(cells)

    if not headers:
        return i

    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(headers):
        if j < len(hdr_cells):
            hdr_cells[j].text = h
            for para in hdr_cells[j].paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < len(row_cells):
                # Remove markdown bold markers
                clean = cell_text.replace("**", "")
                row_cells[c_idx].text = clean
                for para in row_cells[c_idx].paragraphs:
                    for run in para.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11)

    doc.add_paragraph()
    return i


def apply_inline_formatting(para, text):
    """Add a run to paragraph with basic inline bold/italic markdown handling."""
    # Split on ** (bold) and * (italic)
    segments = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for seg in segments:
        if seg.startswith("**") and seg.endswith("**"):
            run = para.add_run(seg[2:-2])
            run.bold = True
        elif seg.startswith("*") and seg.endswith("*") and len(seg) > 2:
            run = para.add_run(seg[1:-1])
            run.italic = True
        elif seg.startswith("`") and seg.endswith("`") and len(seg) > 2:
            run = para.add_run(seg[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            para.add_run(seg)
    for run in para.runs:
        if not run.font.name or run.font.name not in ("Courier New",):
            run.font.name = "Times New Roman"
            if not run.font.size:
                run.font.size = Pt(12)


def add_image_from_md(doc, alt_text, image_ref):
    """Insert a Markdown image into the Word report with page-safe scaling."""
    image_path = Path(image_ref)
    if not image_path.is_absolute():
        image_path = (BASE_DIR / image_path).resolve()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not image_path.exists():
        run = p.add_run(f"[Missing image: {alt_text}]")
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.italic = True
        return

    with Image.open(image_path) as im:
        pixel_width, pixel_height = im.size

    max_width = 6.1
    max_height = 7.2
    aspect = pixel_width / pixel_height if pixel_height else 1
    render_width = min(max_width, max_height * aspect)

    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(render_width))


def build_document():
    doc = Document()
    set_doc_margins(doc)
    configure_styles(doc)
    add_cover_page(doc)

    with open(REPORT_PATH, "r", encoding="utf-8") as f:
        raw = f.read()

    lines = raw.split("\n")
    i = 0
    in_code_block = False
    code_lines = []
    last_was_image = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Code block ──────────────────────────────────────────────
        if stripped.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                in_code_block = False
                # Add code paragraph
                for cl in code_lines:
                    p = doc.add_paragraph(style="Normal")
                    run = p.add_run(cl if cl else " ")
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.left_indent = Cm(1)
                doc.add_paragraph()
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        image_match = re.match(r"^!\[(.*?)\]\((.*?)\)$", stripped)
        if image_match:
            add_image_from_md(doc, image_match.group(1), image_match.group(2))
            last_was_image = True
            i += 1
            continue

        if last_was_image and stripped.startswith("*Figure ") and stripped.endswith("*"):
            p = doc.add_paragraph(style="Normal")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(stripped[1:-1])
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.italic = True
            last_was_image = False
            i += 1
            continue

        # ── Skip raw HTML-style lines ────────────────────────────────
        if stripped:
            last_was_image = False

        if stripped.startswith("<!--") or stripped.startswith("<"):
            i += 1
            continue

        # ── Horizontal rule ──────────────────────────────────────────
        if stripped in ("---", "***", "___"):
            doc.add_paragraph()
            i += 1
            continue

        # ── Page break marker (we use --- for section separation) ────
        # (handled above as horizontal rule / spacing)

        # ── Headings ─────────────────────────────────────────────────
        if stripped.startswith("#### "):
            p = doc.add_paragraph(stripped[5:], style="Heading 3")
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(stripped[4:], style="Heading 3")
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(stripped[3:], style="Heading 2")
            i += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph(stripped[2:], style="Heading 1")
            i += 1
            continue

        # ── Table ────────────────────────────────────────────────────
        if stripped.startswith("|"):
            i = add_table_from_md(doc, lines, i)
            continue

        # ── Blockquote (NOTE TO STUDENT boxes) ───────────────────────
        if stripped.startswith("> "):
            content = stripped[2:]
            # Collect multi-line blockquotes
            bq_lines = [content]
            while i + 1 < len(lines) and lines[i + 1].strip().startswith("> "):
                i += 1
                bq_lines.append(lines[i].strip()[2:])
            full_bq = " ".join(bq_lines)
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.right_indent = Cm(1.5)
            run = p.add_run(full_bq.replace("**", ""))
            run.font.italic = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(80, 80, 80)
            i += 1
            continue

        # ── Bullet list items ─────────────────────────────────────────
        if stripped.startswith("- ") or stripped.startswith("* "):
            content = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            apply_inline_formatting(p, content)
            i += 1
            continue

        # ── Numbered list items ───────────────────────────────────────
        num_match = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if num_match:
            content = num_match.group(2)
            p = doc.add_paragraph(style="List Number")
            apply_inline_formatting(p, content)
            i += 1
            continue

        # ── Empty line → paragraph spacing ───────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Regular paragraph ─────────────────────────────────────────
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        apply_inline_formatting(p, stripped)
        i += 1

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
